def web_scraper_pro(url):
    if 'timesofindia' in url:
        return app(url)
    else:
        return pypeter(url)

def app(url):
    #scrapping part
    import pdfkit


    pdf_file = 'test.pdf'


    options = {
        'page-size': 'Letter',
        'margin-top': '0mm',
        'margin-right': '0mm',
        'margin-bottom': '0mm',
        'margin-left': '0mm',
        'no-print-media-type': None, 
    }



    # Convert the webpage to PDF and save it to disk
    pdfkit.from_url(url, pdf_file, options=options)
    return pdf2extract(pdf_file)

def pdf2extract(pdf_file):
    
    import statistics
    # calculate the mode

    import fitz
    from docx import Document
    from docx2pdf import convert
    from docx import Document


    def finding_modd(c):
        max_length = 0
        max_index = -1

    # Find the max length of the 4th index in the "c" list
        for index, sublist in enumerate(c):
            if len(sublist[4]) > max_length:
                max_length = len(sublist[4])
                max_index = index

        # Extract the value from the 1st index in the corresponding sublist
        if max_index != -1:
            result = c[max_index][0]
        # print(result)
        return result



    def final_news():
        c=[]
        mod_finding = []

        # open the PDF file
        with fitz.open('test.pdf') as doc:

            # iterate over each page in the document
            for page in doc:

                # extract the text blocks from the page
                blocks = page.get_text_blocks()
                
                
            
                for block in blocks:
                    # print(block)
                    c.append(block)
        # print(c)
                # print(c)
        print("hiiii")

        for i in c:
            mod_finding.append(i[0])
        # print(c)
        mod2 = statistics.mode(mod_finding)
        # print(mod2)
        mod = finding_modd(c)
        # print(mod)
        # print(mod)
        print('hi mod is complete')
        news = []
        document = Document()
        # for i in c:
        # headline = []
        for i in c:
            # print(i)
            head = 0
            if i[0] == mod2:
                        if "/" not in i[4] and "\\" not in i[4]  and "image" not in i[4] and "READ" not in i[4] and "Also Read" not in i[4] and not i[4].isupper() and (len(i[4].split()) > 4 or "." in i[4]):
                            print(i)
                            news.append(i[4])
                            head = head+1
                            # if(head == 1):
                                # headline.append(i[4])
                            document.add_paragraph(i[4])
        document.save("test.docx")
        convert("test.docx", "test.pdf")
    final_news()
    return "test.docx"

def pypeter(url):
    from pyppeteer import launch
    import asyncio
    async def generate_pdf(url):
        # Launch a headless browser
        browser = await launch()

        # Open a new tab and navigate to the URL
        page = await browser.newPage()
        await page.goto(url)

        # Generate the PDF
        pdf_bytes = await page.pdf()

        # Close the browser
        await browser.close()

        return pdf_bytes

    # Generate the PDF using Pyppeteer
    pdf_bytes = asyncio.run(generate_pdf(url))

    # Save the PDF to a file
    with open('test.pdf', 'wb') as f:
        f.write(pdf_bytes)
    pdf2extract("test.pdf")
    return pdf3extract()

def pdf3extract():
    import statistics
    # calculate the mode

    import fitz
    from docx import Document
    import PyPDF2


    def final_news():
        c=[]
        mod_finding = []

        # open the PDF file

        # Open the PDF file in read-binary mode
        with open('test.pdf', 'rb') as file:
            # Create a PDF reader object
            reader = PyPDF2.PdfReader(file)

            # Get the total number of pages in the PDF
            num_pages = len(reader.pages)

            # Iterate over each page
            for page_number in range(num_pages):
                # Get the current page
                page =  reader.pages[page_number]

                # Extract the text from the page
                text = page. extract_text()

                # Split the text into lines
                lines = text.split('\n')

                # Process each line
                for line in lines:
                    # Skip empty lines
                    if line.strip() == '':
                        continue

                    # Process the line as needed
                    c.append(line)




        # print(c)
                # print(c)
        print("hiiii")


        news = []
        document = Document()

        headline = []
        for i in c:
                        if "©" not in i:
                            print(i)
                            if "BBC" not in i and  "©" not in i and  "+" not in i and "|" not in i and "/" not in i and "\\" not in i and "image" not in i and "READ" not in i and "Also Read" not in i and not i.isupper() and (len(i.split()) > 4 or "." in i):
                                news.append(i)
                        else:
                            break;
                                # headline.append(i[4])
        document.add_paragraph(news)

        document.save("test.docx")
    final_news()
    return "test.pdf"

def read_final_data():
    import fitz
    from docx import Document
    import PyPDF2

    # Set initial values
    title = ""
    content = ""

    # Open the PDF file
    with open('test.pdf', 'rb') as file:
        # Create a PDF reader object
        reader = PyPDF2.PdfReader(file)

        # Get the total number of pages in the PDF
        num_pages = len(reader.pages)


            # Get the current page
        page = reader.pages[0]

            # Extract the text from the page
        text = page.extract_text()

            # Split the text into lines
        lines = text.split('\n')

            # Process each line
            # for index, line in enumerate(lines):
            #     # Skip empty lines
            #     if line.strip() == '':
            #         continue
                # First line is considered as the title
        title = title+lines[0]
        title = title + lines[1]
        # title = lines[0:2]
        content = lines[1:]
        print("printing lines :")
        print(lines)

    # Return the title and content
    return title, content


link = input("enter website link: \n")
web_scraper_pro(link)

title,content = read_final_data()
print(title)
print(content)
